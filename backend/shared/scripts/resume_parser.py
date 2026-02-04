import sys
import os
import json
import re
import fitz  # PyMuPDF
import docx
from datetime import datetime

def clean_text(text):
    if not text:
        return ""
    # Remove control characters but keep newlines/tabs
    text = "".join(char for char in text if char.isprintable() or char in ["\n", "\t", "\r"])
    # Fix common ligatures that break AI (e.g., "fi" as a single character)
    text = text.replace('\uf0b7', '•').replace('\uf02d', '-')
    # Remove multiple spaces while keeping newlines
    text = re.sub(r' [ ]+', ' ', text)
    return text.strip()

def extract_pdf_advanced(file_path):
    text_blocks = []
    links = set()
    try:
        doc = fitz.open(file_path)
        for page in doc:
            # 1. Extract text with block identification (Handles columns perfectly)
            blocks = page.get_text("blocks")
            # Sort blocks: top-to-bottom, then left-to-right
            blocks.sort(key=lambda b: (b[1], b[0]))
            for b in blocks:
                if b[4].strip():
                    text_blocks.append(b[4].strip())
            
            # 2. Extract links
            page_links = page.get_links()
            for link in page_links:
                if 'uri' in link:
                    uri = link['uri'].strip()
                    if uri.startswith(('http', 'https', 'mailto')):
                        links.add(uri)
        
        doc.close()
    except Exception as e:
        print(f"Error PDF: {str(e)}", file=sys.stderr)
    
    return "\n\n".join(text_blocks), list(links)

def extract_docx_advanced(file_path):
    text = []
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())
        
        # Also get table data (Common in older resumes)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
                    
    except Exception as e:
        print(f"Error DOCX: {str(e)}", file=sys.stderr)
    
    return "\n".join(text), []

def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "No file path provided"}))
        return

    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(json.dumps({"error": "File not found"}))
        return

    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    links = []

    if ext == ".pdf":
        text, links = extract_pdf_advanced(file_path)
    elif ext in [".docx", ".doc"]:
        text, links = extract_docx_advanced(file_path)
    else:
        print(json.dumps({"error": f"Unsupported format: {ext}"}))
        return

    cleaned_text = clean_text(text)
    
    # Quick sanity check for contact info (Internal validation)
    emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', cleaned_text)
    
    result = {
        "text": cleaned_text,
        "links": links,
        "emails": list(set(emails)),
        "length": len(cleaned_text),
        "timestamp": datetime.now().isoformat()
    }

    print(json.dumps(result))

if __name__ == "__main__":
    main()
